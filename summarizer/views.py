import json
from docx import Document
from django.http import HttpResponse
from django.views.decorators.http import require_http_methods
from .utils import extract_text_from_txt, extract_text_from_pdf, extract_text_from_url, generate_summary, \
    sanitize_input, validate_uploaded_file, extract_and_validate_url, extract_text_from_docx, extract_text_from_audio
from rest_framework.decorators import api_view
from rest_framework.response import Response
from .serializers import SummarizationSerializer
from django.views.decorators.csrf import csrf_protect, csrf_exempt
from django.core.exceptions import ValidationError

MAX_TEXT_LENGTH = 4096


@api_view(['POST'])
def summarize_text(request):
    """
    This view handles the summarization of text based on the input type.
    It supports text, file (txt/pdf), and URL inputs.

    Args:
        request (HttpRequest): The HTTP request object containing the data.

    Returns:
        JsonResponse: A response with either the summary or an error message.
    """
    serializer = SummarizationSerializer(data=request.data)

    if serializer.is_valid():
        data = serializer.validated_data
        input_type = data['input_type']
        form = data['form']
        length = data['length']
        language = data['language']
        granularity = data['granularity']
        text = None

        if input_type == 'text':
            raw_text = data['text']
            url = extract_and_validate_url(raw_text)

            if url and raw_text.startswith(url):
                text = extract_text_from_url(url)
                if not text:
                    return Response({'error': 'Invalid or unsafe URL'}, status=400)
            else:
                text = sanitize_input(raw_text)
                if len(text) > MAX_TEXT_LENGTH:
                    return Response({'error': f'Text is too long! Max size is {MAX_TEXT_LENGTH} characters.'}, status=400)

        elif input_type == 'file':
            uploaded_file = data['file']

            try:
                validate_uploaded_file(uploaded_file)
            except ValidationError as e:
                return Response({'error': str(e)}, status=400)

            if uploaded_file.name.endswith('.txt'):
                text = extract_text_from_txt(uploaded_file)
            elif uploaded_file.name.endswith('.pdf'):
                text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.name.endswith('.docx'):
                text = extract_text_from_docx(uploaded_file)
            elif uploaded_file.name.endswith('.wav'):
                text = extract_text_from_audio(uploaded_file)
            else:
                return Response({'error': 'Unsupported file format'}, status=400)

        if not text:
            return Response({'error': 'No valid text found'}, status=400)

        summary = generate_summary(form=form, length=length, language=language, text=text, granularity=granularity)
        return Response({'summary': summary}, status=200)

    return Response({'error': serializer.errors}, status=400)

@csrf_exempt
@require_http_methods(["POST"])
def download_summary(request):

    data = json.loads(request.body)
    summary = data.get('summary', '')

    doc = Document()
    doc.add_heading('Summary Document', 0)
    doc.add_paragraph(summary)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=summary.docx'

    doc.save(response)

    return response